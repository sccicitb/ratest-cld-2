"""Tests for app.rag.extract (Task 3.2 / Task 2 PDFOxide swap): PDFOxide/docx/text routing.

Routing tests monkeypatch `pdf_oxide.PdfDocument` methods so they're fast and
never load the real PaddleOCR models.
"""
from __future__ import annotations

import json

import pytest

from app.rag.extract import SUPPORTED_KB_TYPES, extract_text

OCR_SENTINEL = "OCR_TEXT"


@pytest.fixture()
def blob_dir(tmp_path, monkeypatch):
    from app.config import settings

    monkeypatch.setattr(settings, "blob_dir", str(tmp_path))
    return tmp_path


def _write(blob_dir, name: str, content: bytes) -> str:
    path = blob_dir / name
    path.write_bytes(content)
    return name


def test_supported_kb_types_contains_expected_extensions():
    assert SUPPORTED_KB_TYPES == {".pdf", ".md", ".txt", ".docx", ".doc", ".csv", ".json"}


def test_extract_text_layer_pdf_uses_native_text(blob_dir):
    """A real text PDF extracts its text via PDFOxide (no OCR needed)."""
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Hello from a text layer.")
    key = _write(blob_dir, "text.pdf", doc.write())
    doc.close()

    out = extract_text(key, "text.pdf")
    assert "Hello from a text layer." in out


def test_extract_pdf_routes_to_ocr_when_enabled(blob_dir, monkeypatch):
    """ocr_enabled=True → each page read via extract_text_auto (OCR-capable)."""
    import fitz
    import pdf_oxide

    from app.config import settings
    monkeypatch.setattr(settings, "ocr_enabled", True)

    doc = fitz.open()
    doc.new_page()
    key = _write(blob_dir, "scan.pdf", doc.write())
    doc.close()

    monkeypatch.setattr(
        pdf_oxide.PdfDocument, "extract_text_auto",
        lambda self, page: OCR_SENTINEL, raising=True,
    )
    out = extract_text(key, "scan.pdf")
    assert OCR_SENTINEL in out


def test_extract_pdf_skips_ocr_when_disabled(blob_dir, monkeypatch):
    """ocr_enabled=False → native extract_text only; extract_text_auto NOT called."""
    import fitz
    import pdf_oxide

    from app.config import settings
    monkeypatch.setattr(settings, "ocr_enabled", False)

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "native only")
    key = _write(blob_dir, "native.pdf", doc.write())
    doc.close()

    def _boom(self, page):
        raise AssertionError("extract_text_auto must not be called when ocr_enabled=False")

    monkeypatch.setattr(pdf_oxide.PdfDocument, "extract_text_auto", _boom, raising=True)
    out = extract_text(key, "native.pdf")
    assert "native only" in out


def test_open_pdf_oxide_repairs_malformed_pdf(blob_dir, monkeypatch):
    """A PDF PDFOxide can't open is repaired via PyMuPDF, then opened."""
    import fitz
    import pdf_oxide

    from app.rag.extract import _open_pdf_oxide

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "repair me")
    data = doc.write()
    doc.close()

    real_from_bytes = pdf_oxide.PdfDocument.from_bytes
    calls = {"n": 0}

    def flaky_from_bytes(b):
        calls["n"] += 1
        if calls["n"] == 1:
            raise OSError("Invalid cross-reference table")
        return real_from_bytes(b)

    monkeypatch.setattr(
        pdf_oxide.PdfDocument, "from_bytes", staticmethod(flaky_from_bytes), raising=True
    )
    pdoc = _open_pdf_oxide(data)
    assert calls["n"] == 2  # first raised, repaired bytes opened on retry
    assert "repair me" in pdoc.extract_text(0)


def test_extract_text_from_docx(blob_dir):
    import docx

    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    docx_path = blob_dir / "sample.docx"
    document.save(str(docx_path))

    text = extract_text("sample.docx", "sample.docx")
    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_extract_text_from_txt(blob_dir):
    _write(blob_dir, "notes.txt", "plain text content".encode("utf-8"))
    assert extract_text("notes.txt", "notes.txt") == "plain text content"


def test_extract_text_from_md(blob_dir):
    _write(blob_dir, "readme.md", "# Heading\nbody".encode("utf-8"))
    assert extract_text("readme.md", "readme.md") == "# Heading\nbody"


def test_extract_text_from_csv(blob_dir):
    _write(blob_dir, "data.csv", "a,b\n1,2".encode("utf-8"))
    assert extract_text("data.csv", "data.csv") == "a,b\n1,2"


def test_extract_text_from_json(blob_dir):
    payload = json.dumps({"k": "v"})
    _write(blob_dir, "data.json", payload.encode("utf-8"))
    assert extract_text("data.json", "data.json") == payload


def test_extract_text_legacy_doc_raises_clear_error(blob_dir):
    _write(blob_dir, "legacy.doc", b"\xd0\xcf\x11\xe0fake-ole-header")
    with pytest.raises(Exception) as exc_info:
        extract_text("legacy.doc", "legacy.doc")
    assert "libreoffice" in str(exc_info.value).lower() or "not" in str(exc_info.value).lower()


def test_extract_text_unknown_extension_raises_value_error(blob_dir):
    _write(blob_dir, "image.png", b"\x89PNG fake")
    with pytest.raises(ValueError):
        extract_text("image.png", "image.png")
